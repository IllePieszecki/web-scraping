
import os
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Inches

from selenium.webdriver.common.by import By
from components.env import Config

from components.driver_browser import Browser
from components.web_page import WebPage
from datetime import date


class Buyee(WebPage, Browser, Config):
    def __init__(self):
        # Inicializa la clase base
        self.initialize_browser()
        self.load_environment_variables()
        WebPage.__init__(self, self.driver)

    def login(self):
        self.driver.get("https://buyee.jp/")
        self.wait_page(2)
        if self.check_if_exists(".g-modalBanner__closeButton.blue"):
            self.find(".g-modalBanner__closeButton.blue").click()
        self.find("a.button.link").click()

        self.find("#login_mailAddress").send_keys(os.getenv('BUYEE_USER'))
        self.find("#login_password").send_keys(os.getenv('BUYEE_PASSWD'))
        self.find(".btn_signup_blue").click()

        self.wait_page(2)

    def extract_products(self):
        self.driver.get("https://buyee.jp/mybaggages/shipped/1")
        self.wait_for_element("#luggageInfo_collection")

        self.__delete_image_files(os.getenv('IMAGE_PATH', '.'))

        first_order = self.find('#luggageInfo_collection li')
        first_order.find_element(By.CSS_SELECTOR, ".g-feather-chevron-up").click()

        order_table = first_order.find_element(By.CSS_SELECTOR, "table.luggageInfo_order")
        products = order_table.find_elements('css selector', 'tbody tr:nth-child(n+2)')
        # print(len(products))

        for product in products:
            columns = product.find_elements(By.TAG_NAME, 'td')
            row_data = [col.text for col in columns]
            page_origin = row_data[0]
            order_num = row_data[1]
            link = columns[2].find_element(By.TAG_NAME, 'a').get_attribute('href')

            match page_origin:
                case 'mercari':
                    self._mercari(link, order_num)
                case 'Yahoo! Japan Auction':
                    self._yahoo(link, order_num)
                case 'JDirectItems Auction':
                    self._yahoo(link, order_num)
                case 'PayPay Flea market':
                    pass
                case 'JYP JAPAN ONLINE STORE':
                    self._jyp(link, order_num)
                case _:
                    print(f'La página de origen "{page_origin}" no esta programada...')

    def _mercari(self, link: str, order_num: str):
        self.__change_tab(link, '#content_wrap', '.m-goodsButton', order_num)

    def _yahoo(self, link: str, order_num: str):
        self.__change_tab(link, '#content_wrap', '#auction_item_description', order_num)

    def _flea_market(self, link: str):
        pass

    def _jyp(self, link: str, order_num: str):
        self.__change_tab(link, '#breadcrumb', 'h2.section-heading span', order_num)

    def __change_tab(self, href_value: str, header: str, footer: str, order_num: str):

        self.driver.execute_script(f"window.open('{href_value}', '_blank');")
        self.driver.switch_to.window(self.driver.window_handles[1])
        header = self.find(header)
        footer = self.find(footer)

        x_left = header.rect['x']
        x_right = header.rect['width'] + x_left + 30
        y_right = footer.rect['y'] - header.rect['y']

        self.driver.execute_script("arguments[0].scrollIntoView();", header)
        self.wait_page(2)
        self.__take_screenshot(order_num + '.jpg', x_left, x_right, y_right)
        self.wait_page(3)
        self.driver.close()  # Cierra la ventana actual
        self.driver.switch_to.window(self.driver.window_handles[0])

    def __delete_image_files(self, image_path: str):
        for filename in os.listdir(image_path):
            if filename.endswith(".png") or filename.endswith(".jpg"):
                file_path = os.path.join(image_path, filename)
                try:
                    os.remove(file_path)
                except Exception as e:
                    print(f"Error al eliminar el archivo {file_path}: {e}")

    def __take_screenshot(self, filename: str, x_left: float, x_right: float, y_right: float):
        # print(f"{x_left},0,{x_right},{y_right}")
        screenshot_path = os.getenv('IMAGE_PATH', '.')
        full_path = os.path.join(screenshot_path, filename)
        screenshot = self.driver.get_screenshot_as_png()
        screenshot = Image.open(BytesIO(screenshot))
        #                                     left_arriba / arriba / right_arriba / abajo
        cropped_screenshot = screenshot.crop((x_left, 0, x_right, y_right))
        cropped_screenshot.save(full_path)

    def create_doc(self):
        today = date.today()
        doc_name = 'Buyee_' + today.strftime("%d_%m_%Y") + '.docx'
        doc = Document()
        doc_path = os.getenv('IMAGE_PATH', '.')
        full_path = os.path.join(doc_path, doc_name)

        for filename in os.listdir(doc_path):
            if filename.endswith(".png") or filename.endswith(".jpg") or filename.endswith(".jpeg"):
                # Ruta completa de la imagen
                image_path = os.path.join(doc_path, filename)
                # Insertar imagen en el documento
                doc.add_picture(image_path, width=Inches(7))  # Ajustar el ancho de la imagen
                # Agregar un espacio después de la imagen
                doc.add_paragraph("\n")

        doc.save(full_path)
        print(f"Documento guardado como: {full_path}")

    def order_extract(self):
        self.login()
        self.extract_products()
        self.close()
        self.create_doc()

    def close(self):
        self.driver.quit()


buyee = Buyee()  # Inicializa Buyee
buyee.order_extract()
